#!/usr/bin/env python3
"""
ARC-AGI Testing Report Generator
Collects all .go files and all .json result files into a DOCX report.
"""

import os
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt


def read_file_content(file_path: Path) -> str:
    """Read and return file content."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Error reading file: {e}"


def format_json_content(json_path: Path) -> str:
    """Read and pretty-format JSON content."""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return json.dumps(data, indent=2)
    except Exception as e:
        return f"Error parsing JSON: {e}"


def create_report(go_files: list, json_files: list, output_path: str):
    """Generate the DOCX report."""
    doc = Document()
    
    # Title
    doc.add_heading('ARC-AGI Testing Report', 0)
    doc.add_paragraph(f'Go files: {len(go_files)} | JSON files: {len(json_files)}')
    doc.add_paragraph('=' * 50)
    
    # Add all Go files
    doc.add_heading('Go Source Files', level=1)
    
    for go_file in go_files:
        doc.add_heading(go_file.name, level=2)
        go_content = read_file_content(go_file)
        
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(go_content[:50000])
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(8)
        
        if len(go_content) > 50000:
            doc.add_paragraph(f"... [truncated, {len(go_content)} total characters]")
        
        doc.add_page_break()
    
    # Add all JSON files
    doc.add_heading('JSON Result Files', level=1)
    
    for json_file in json_files:
        doc.add_heading(json_file.name, level=2)
        json_content = format_json_content(json_file)
        
        json_para = doc.add_paragraph()
        json_run = json_para.add_run(json_content[:30000])
        json_run.font.name = 'Courier New'
        json_run.font.size = Pt(8)
        
        if len(json_content) > 30000:
            doc.add_paragraph(f"... [truncated, {len(json_content)} total characters]")
        
        doc.add_page_break()
    
    doc.save(output_path)
    print(f"✅ Report saved to: {output_path}")


def main():
    script_dir = Path(__file__).parent.resolve()
    print(f"📁 Scanning directory: {script_dir}")
    
    # Find all files
    go_files = sorted(script_dir.glob("*.go"))
    json_files = sorted(script_dir.glob("*.json"))
    
    print(f"📄 Found {len(go_files)} Go files")
    print(f"📊 Found {len(json_files)} JSON files")
    
    # Generate report
    output_path = script_dir / "arcagi_testing_report.docx"
    create_report(go_files, json_files, str(output_path))
    
    print("\n📋 Go files:")
    for f in go_files:
        print(f"  ✓ {f.name}")
    
    print("\n📋 JSON files:")
    for f in json_files:
        print(f"  ✓ {f.name}")


if __name__ == "__main__":
    main()
